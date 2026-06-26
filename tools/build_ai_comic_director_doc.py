from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("AI漫剧导演.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x1F, 0x1F)
MUTED = RGBColor(0x66, 0x66, 0x66)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"


def set_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, before=0, after=6, line_spacing=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def add_para(doc, text="", style=None, before=0, after=6, line_spacing=1.25, bold_lead=None):
    p = doc.add_paragraph(style=style)
    style_paragraph(p, before=before, after=after, line_spacing=line_spacing)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_font(r, color=INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = add_para(doc, item, style="List Bullet", after=4)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)


def add_numbers(doc, items):
    for item in items:
        p = add_para(doc, item, style="List Number", after=4)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    style_paragraph(
        p,
        before={1: 18, 2: 14, 3: 10}.get(level, 8),
        after={1: 10, 2: 7, 3: 5}.get(level, 5),
        line_spacing=1.15,
    )
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_label_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    widths = [1700, 7660]
    for i, heading in enumerate(("模块", "执行要求")):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, HEADER_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(heading)
        set_font(run, bold=True, color=DARK_BLUE)
    for label, detail in rows:
        cells = table.add_row().cells
        for i, text in enumerate((label, detail)):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_font(run, bold=(i == 0), color=INK)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    style_paragraph(p, before=2, after=2, line_spacing=1.2)
    r1 = p.add_run(title + "：")
    set_font(r1, bold=True, color=DARK_BLUE)
    r2 = p.add_run(body)
    set_font(r2, color=INK)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("AI漫剧导演工作手册")
    set_font(hr, size=9, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(title, after=4, line_spacing=1.1)
    r = title.add_run("AI漫剧导演工作手册")
    set_font(r, size=24, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=12, line_spacing=1.2)
    r = subtitle.add_run("面向 AI 视频、AI 生图、短剧分镜与提示词工程的项目级生产规范")
    set_font(r, size=11, color=MUTED)

    add_callout(
        doc,
        "使用定位",
        "这份文档把原始“虚拟导演提示词”整理为可执行的生产手册：先保护剧本，再建立空间与角色连续性，最后输出可直接交给 AI 视频模型使用的分镜提示词、九宫格分镜和质量评分。",
    )

    add_heading(doc, "1. 项目目标", 1)
    add_para(
        doc,
        "本项目的目标不是写普通提示词，而是建立一套可复用的 AI漫剧生产系统。系统必须同时承担虚拟导演、分镜导演、摄影指导、视效总监和 AI 视频提示词工程师的职责。",
    )
    add_bullets(
        doc,
        [
            "把剧本片段拆成可生成的 15 秒独立视频段落。",
            "把每段补齐人物、空间、动作、光影、声音、技术锁和负面约束。",
            "让每段视频即使独立生成，也能保持角色一致、站位稳定、动作连续。",
            "为九宫格分镜、关键帧图和质量评分提供统一标准。",
        ]
    )

    add_heading(doc, "2. 不可破坏的核心规则", 1)
    add_label_table(
        doc,
        [
            ("剧本忠诚", "用户提供的台词一字不改，不删除、不重排、不添加。长台词可以通过反应镜头拆开，但台词文本和顺序必须保持。"),
            ("段落独立", "每个 15 秒视频段都必须重新交代人物身份、外貌、站位、场景锚点、START_POSE、END_POSE、光影和动作关系。"),
            ("屏幕坐标", "当使用参考图、站位图、九宫格或“以画面为参照”时，所有方向以观众屏幕为准：画面左、画面右、前景、中景、后景。"),
            ("角色控制", "每段重点人物最多 3-4 个。其他人物只做虚化背影、肩膀、手部、反光轮廓、离焦人影或背景形状。"),
            ("物理连续", "动作要符合方向、力、惯性、反作用和落点。禁止瞬移、换位、穿模、肢体融合、空间漂移和左右翻轴。"),
            ("完整输出", "用户要求“最终版”“最高评分版”“9.6+”时，必须输出完整成品，不能写“同上”或用概括替代正文。"),
        ]
    )

    add_heading(doc, "3. 输入处理流程", 1)
    add_numbers(
        doc,
        [
            "锁定不可变输入：台词、角色名、剧情事件、指定风格、时长、参考图、平台限制和安全要求。",
            "提取人物表：每个角色写清年龄段、外貌、服装、材质、状态、情绪和动作倾向。",
            "建立空间图：谁在画面左/右/前/后，谁面对谁，关键道具、门、墙、屏幕、裂缝、车辆、武器、高台在哪里。",
            "判断拆分粒度：一个约 15 秒剧情切片默认拆成 5-8 个电影级镜头；复杂动作或多段台词应拆成更多段。",
            "确定风格优先级：本轮明确指定风格 > 项目已定风格 > 默认 3D动漫质感 / UE5 / PBR / 电影级虚拟摄影。",
        ]
    )

    add_heading(doc, "4. 15 秒分镜提示词标准模板", 1)
    add_para(doc, "每个独立视频段建议使用以下模块。模块可以按剧情增减细节，但不能省略连续性、互动锁和负面约束。")
    add_label_table(
        doc,
        [
            ("[分镜编号]", "XX｜标题｜最终版。标题要直接指向冲突或情绪爆点。"),
            ("黄金前3秒", "一句话概括最抓人的核心画面，必须包含冲突、动作或情绪钩子。"),
            ("[人物]", "列出重点人物及外貌、服装、状态、情绪。背景人物只做虚化或局部轮廓。"),
            ("[人物详细站位]", "明确屏幕坐标、START_POSE、动作坐标、接触点、END_POSE。"),
            ("[场景空间]", "写 LOC、空间结构、空气层、环境情绪、时代锚点、渲染风格。"),
            ("[人物、动作与台词]", "按 5-8 个镜头写焦段、机位、景别、运镜、时间段、动作起点/过程/落点和原文台词。"),
            ("[镜头语言]", "解释情绪递进、动作承接、视线引导、反应镜头和结尾钩子。"),
            ("[视觉特效/UI]", "写粒子、体积雾、能量、屏幕 UI、镜面反射等。UI 只允许存在于设备屏幕内部。"),
            ("[光影氛围]", "写主光、辅光、轮廓光、负补光、动态光源、色彩情绪和必要色温。"),
            ("[声音设计]", "写环境声、动作声、设备/特效声、对白声线和音乐要求。默认无背景音乐。"),
            ("[互动锁定]", "锁人物位置、距离、朝向、接触点、道具位置和动作顺序。"),
            ("[连续性技术锁]", "锁造型、场景锚点、摄影轴线、动作方向、光色、台词顺序和物理规则。"),
            ("[关键帧]", "KEYFRAME_START / KEYFRAME_MID / KEYFRAME_END。"),
            ("[负面约束]", "无外挂字幕、水印、logo、旁白文字、多余人物、换位、翻轴、瞬移、穿模、肢体融合、低清、AI感、画面外 UI。"),
            ("[节奏描述]", "前3秒强钩子，中段加速，后段落点，结尾留下一段钩子。"),
            ("[风格]", "优先使用用户指定风格；未指定时使用项目默认风格。"),
        ]
    )

    add_heading(doc, "5. 单镜头写法要求", 1)
    add_para(doc, "每个镜头都要像可执行拍摄指令，而不是概念描述。推荐格式如下：")
    add_callout(
        doc,
        "镜头格式",
        "【焦段mm / 机位 / 景别｜专业运镜术语｜时间段】当前站位 + 动作起点 + 动作过程 + 动作落点 + 前景遮挡 + 表情变化 + 视线方向 + 与前后镜头的动作承接 + 光影落点。",
    )
    add_bullets(
        doc,
        [
            "焦段示例：24mm 建立空间压迫，35mm 跟随动作，50mm 角色关系，85mm 情绪特写，100mm 微距道具。",
            "机位示例：低角度压迫、肩后过肩、侧后方跟拍、贴地仰拍、俯视俯压、前景遮挡窥视。",
            "运镜示例：Steadicam、Dolly Track、Technocrane、Crash Zoom、Match Cut、探针微距、动接动剪辑。",
            "表演示例：气口、停顿、重音、句尾落点、呼吸节奏、眼神变化、微表情和情绪层次。",
        ]
    )

    add_heading(doc, "6. 九宫格分镜模板", 1)
    add_para(doc, "当用户要求“九宫格”“分镜图”“等比例多宫格”“站位图”时，使用九个独立关键镜头建立空间和动作连续性。")
    add_label_table(
        doc,
        [
            ("规格", "3x3 九宫格；所有画格等比例；单格建议 16:9；画面方位必须和视频提示词一致。"),
            ("每格字段", "九宫格编号、标题、景别、画面、站位、重点。"),
            ("统一约束", "同一空间坐标、同一摄影轴线、不翻轴、角色初始位置固定、关键道具固定、动作可连续剪辑。"),
            ("清洁画面", "如需干净画面，加入：整体画面干净、平滑、统一，大色块叙事，边缘清晰利落，无高频噪点。"),
        ]
    )

    add_heading(doc, "7. 光影、声音与风格默认值", 1)
    add_bullets(
        doc,
        [
            "默认视觉：3D动漫质感 / 3D游戏CG / UE5渲染 / PBR材质 / 电影级虚拟摄影 / 好莱坞级电影视效 / ASC电影摄影标准。",
            "真人写实时：真人写实风格 / 好莱坞级电影视效 / ASC电影摄影标准 / 高级实拍质感。",
            "科幻光色示例：冷蓝应急灯约 6800K，红色警报光约 2300K，屏幕蓝光约 7200K，火花光约 3800K-4200K，裂缝冷白光约 7000K。",
            "声音默认：无背景音乐，除非用户另行指定；必须写清环境声、动作声、设备声和对白声线。",
        ]
    )

    add_heading(doc, "8. 安全过审表达", 1)
    add_para(doc, "暗黑、修仙、战斗、恐怖或高压场景中，要保留张力但降低高风险直白词。")
    add_label_table(
        doc,
        [
            ("高风险直白词", "可替代表达"),
            ("恶鬼", "深渊巨影、黑暗轮廓、不可名状的压迫性存在"),
            ("地狱", "绝境领域、封锁危险区、深层禁区"),
            ("杀戮", "毁灭性压迫、致命冲击、失控打击"),
            ("血腥", "暗红能量痕迹、破碎红光、红色冲击残影"),
            ("尸体", "倒伏身影、失去行动能力的轮廓、坠落剪影"),
        ]
    )

    add_heading(doc, "9. 质量评分卡", 1)
    add_para(doc, "评分必须诚实。低于 9.6 分时，要指出具体问题并给出可执行优化方案。")
    add_numbers(
        doc,
        [
            "剧本是否忠实：台词完整、不改字、不乱序。",
            "段落是否独立：人物、空间、动作、光影是否重新建立。",
            "镜头是否能剪：动作、视线、情绪是否连续。",
            "站位是否稳定：是否存在换位、瞬移、空间漂移或翻轴。",
            "物理是否可信：力、惯性、反作用和落点是否明确。",
            "情绪是否递进：气口、停顿、重音、眼神和微表情是否服务剧情。",
            "光影是否高级：主光、辅光、轮廓光、负补光、动态光源和色温是否连贯。",
            "AI 风险是否锁住：负面约束是否覆盖穿模、融合、畸形、字幕、水印、画面外 UI。",
            "时长是否适配：是否需要拆成更多 15 秒段落。",
            "是否达到 9.6+：若没有，说明必须补齐的模块。",
        ]
    )

    add_heading(doc, "10. 项目 Skill 使用方式", 1)
    add_para(
        doc,
        "本仓库已配套生成项目级 Skill：skills/ai-comic-drama-production。后续需要让 Codex 直接使用时，可把该目录复制或安装到 Codex skills 目录，也可以在项目内作为提示词生产规范引用。",
    )
    add_label_table(
        doc,
        [
            ("SKILL.md", "触发条件、核心规则、生产流程、默认风格和输出纪律。"),
            ("references/prompt-contract.md", "完整 15 秒视频段、必备技术模块和九宫格分镜输出契约。"),
            ("references/qa-scorecard.md", "9.6+ 评分维度、诊断格式和可执行优化要求。"),
        ]
    )

    add_heading(doc, "11. 最终交付检查清单", 1)
    add_bullets(
        doc,
        [
            "台词原文完整，顺序正确。",
            "每段都有 START_POSE、动作坐标和 END_POSE。",
            "每段 5-8 个镜头，镜头之间能动接动。",
            "屏幕坐标明确，无角色左右歧义。",
            "重点人物不超过 3-4 个。",
            "光影、声音、VFX、UI、互动锁、连续性锁、关键帧、负面约束完整。",
            "结尾留下下一段可接续的眼神、动作、道具、能量或空间压迫钩子。",
            "评分低于 9.6 时，必须给出具体可执行修改，而不是泛泛评价。",
        ]
    )

    doc.core_properties.title = "AI漫剧导演工作手册"
    doc.core_properties.subject = "AI漫剧提示词、分镜、连续性锁和质量评分规范"
    doc.core_properties.keywords = "AI漫剧, AI视频, 分镜, 提示词, Codex Skill"
    doc.save(OUT)


if __name__ == "__main__":
    build()
